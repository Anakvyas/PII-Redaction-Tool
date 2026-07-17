from __future__ import annotations

import shutil

import pytest

docx = pytest.importorskip("docx")

from replacement.docx_redactor import redact_docx
from schemas.common import PIIEntity, PIIType, RedactionStrategy, TextSpan
from services.extraction.docx_extractor import DocxExtractor


def _entity(entity_id: str, text: str, start: int, pii_type: PIIType = PIIType.PERSON) -> PIIEntity:
    return PIIEntity(
        id=entity_id,
        pii_type=pii_type,
        span=TextSpan(start=start, end=start + len(text)),
        raw_value=text,
        confidence=0.99,
        source_detector="test",
    )


def test_docx_extraction_includes_headers_footers_tables_and_redacts_with_formatting(tmp_path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    document = docx.Document()
    document.sections[0].header.paragraphs[0].add_run("Header Alice")
    body = document.add_paragraph()
    body.add_run("Hello ")
    bold = body.add_run("Alice")
    bold.bold = True
    body.add_run(" and Alice")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Cell Alice"
    document.sections[0].footer.paragraphs[0].add_run("Footer Alice")
    document.save(source)

    extracted = DocxExtractor().extract(str(source), "doc-1")
    text = extracted.flattened_text()
    starts = [i for i in range(len(text)) if text.startswith("Alice", i)]
    entities = [_entity(f"e{i}", "Alice", start) for i, start in enumerate(starts)]

    result = redact_docx(
        str(source),
        str(output),
        entities,
        {PIIType.PERSON: RedactionStrategy.MASK},
    )

    redacted = docx.Document(output)
    flattened = DocxExtractor().extract(str(output), "doc-1").flattened_text()
    body_runs = redacted.paragraphs[0].runs

    assert result.counts_by_type == {PIIType.PERSON: 5}
    assert "Alice" not in flattened
    assert flattened.count("[REDACTED-PERSON]") == 5
    assert body_runs[1].bold is True


def test_docx_extraction_captures_every_cell_in_a_multi_row_multi_column_table(tmp_path) -> None:
    """Regression test for a real, high-severity bug found against an actual
    200+ page document: a Board of Directors table (Name/Designation/DIN/
    Address, 8 rows) extracted only the Name column — every other cell in
    every data row silently vanished.

    Root cause: `_iter_table_cells`'s dedup (for genuinely merged cells)
    tracked `id(cell._tc)` instead of the `_tc` object itself. `row.cells`
    builds a fresh `_Cell` wrapper on each access; keeping only the id
    let CPython recycle a just-freed wrapper's memory address for the
    *next* cell's wrapper, making unrelated cells compare as duplicates.
    A small 1x1 table (the only kind covered elsewhere in this file) never
    has enough live/dead object churn to reproduce it — this table is sized
    to reliably trigger the same object-identity recycling in real CPython."""
    source = tmp_path / "source.docx"

    document = docx.Document()
    table = document.add_table(rows=5, cols=4)
    expected_cells = []
    for row in range(5):
        for col in range(4):
            value = f"R{row}C{col}-{'xyzw'[col]}"
            table.cell(row, col).text = value
            expected_cells.append(value)
    document.save(source)

    text = DocxExtractor().extract(str(source), "doc-1").flattened_text()

    missing = [value for value in expected_cells if value not in text]
    assert missing == [], f"cells silently dropped from extraction: {missing}"


def test_docx_redactor_does_not_replace_partial_words(tmp_path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    document = docx.Document()
    document.add_paragraph("Ann read the Annual Report.")
    document.save(source)

    extracted = DocxExtractor().extract(str(source), "doc-1")
    text = extracted.flattened_text()
    partial_start = text.index("Annual")

    result = redact_docx(
        str(source),
        str(output),
        [_entity("partial", "Ann", partial_start)],
        {PIIType.PERSON: RedactionStrategy.MASK},
    )

    flattened = DocxExtractor().extract(str(output), "doc-1").flattened_text()

    assert result.counts_by_type == {}
    assert "Annual" in flattened
    assert "[REDACTED-PERSON]" not in flattened


def test_docx_redactor_redacts_pii_in_an_embedded_image(tmp_path) -> None:
    """End-to-end: a scanned ID card (or any PII-bearing picture) pasted
    into a DOCX gets OCR'd, the same detection pipeline everything else
    uses finds the PII in it, and the matching regions are blacked out in
    the image itself — not just PII in the document's own text runs."""
    if shutil.which("tesseract") is None:
        pytest.skip("tesseract binary not installed")
    pytest.importorskip("PIL")
    pytest.importorskip("pytesseract")
    from PIL import Image, ImageDraw, ImageFont

    from docx.parts.image import ImagePart

    from services.image_pii_service import ocr_image

    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    font = ImageFont.load_default(size=28)
    image = Image.new("RGB", (800, 200), color="white")
    draw = ImageDraw.Draw(image)
    draw.text((20, 20), "Name: Rashi Patil", fill="black", font=font)
    draw.text((20, 80), "Email: rashi.patil@gmail.com", fill="black", font=font)
    image_path = tmp_path / "id_card.png"
    image.save(image_path)

    document = docx.Document()
    document.add_paragraph("Application attached below.")
    document.add_picture(str(image_path))
    document.save(source)

    result = redact_docx(
        str(source),
        str(output),
        entities=[],  # nothing in the document's own text — only the image has PII
        strategy_map={PIIType.PERSON: RedactionStrategy.MASK, PIIType.EMAIL: RedactionStrategy.MASK},
    )

    assert result.counts_by_type.get(PIIType.EMAIL, 0) >= 1
    assert any(entry.image_filename is not None for entry in result.audit_entries)

    redacted_doc = docx.Document(output)
    image_parts = [p for p in redacted_doc.part.related_parts.values() if isinstance(p, ImagePart)]
    assert len(image_parts) == 1

    redacted_text, _words = ocr_image(image_parts[0].blob)
    assert "rashi.patil@gmail.com" not in redacted_text
